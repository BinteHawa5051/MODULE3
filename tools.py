from typing import List, Dict, Any, Optional
import requests
from datetime import datetime
import json
from io import BytesIO
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import openpyxl


class WebSearchTool:
    """Web search with shallow and deep retrieval modes"""
    
    def __init__(self, api_key: str):
        self.api_key = api_key
        self.base_url = "https://google.serper.dev/search"
    
    def search(self, query: str, depth: str = "shallow") -> Dict[str, Any]:
        """
        Search the web with configurable depth
        depth: 'shallow' (top 3 results) or 'deep' (top 10 results)
        """
        if not self.api_key:
            return {"error": "Serper API key not provided"}
        
        num_results = 3 if depth == "shallow" else 10
        
        payload = json.dumps({
            "q": query,
            "num": num_results
        })
        
        headers = {
            'X-API-KEY': self.api_key,
            'Content-Type': 'application/json'
        }
        
        try:
            response = requests.post(self.base_url, headers=headers, data=payload, timeout=10)
            response.raise_for_status()
            data = response.json()
            
            results = []
            if "organic" in data:
                for item in data["organic"]:
                    results.append({
                        "title": item.get("title", ""),
                        "link": item.get("link", ""),
                        "snippet": item.get("snippet", "")
                    })
            
            return {
                "query": query,
                "depth": depth,
                "results": results,
                "count": len(results)
            }
        except Exception as e:
            return {"error": f"Search failed: {str(e)}"}


class DocumentCreationTool:
    """Create documents in various formats"""
    
    @staticmethod
    def create_pdf(content: str, filename: str) -> BytesIO:
        """Create a PDF document"""
        buffer = BytesIO()
        c = canvas.Canvas(buffer, pagesize=letter)
        width, height = letter
        
        # Simple text wrapping
        y = height - 50
        for line in content.split('\n'):
            if y < 50:
                c.showPage()
                y = height - 50
            c.drawString(50, y, line[:80])  # Limit line length
            y -= 15
        
        c.save()
        buffer.seek(0)
        return buffer
    
    @staticmethod
    def create_word(content: str, filename: str) -> BytesIO:
        """Create a Word document"""
        doc = Document()
        doc.add_heading('Generated Document', 0)
        doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        doc.add_paragraph('')
        
        for paragraph in content.split('\n\n'):
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())
        
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    
    @staticmethod
    def create_excel(data: List[List[Any]], filename: str) -> BytesIO:
        """Create an Excel spreadsheet"""
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Data"
        
        for row in data:
            ws.append(row)
        
        buffer = BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        return buffer


class DateTimeTool:
    """Provides current date and time information"""
    
    @staticmethod
    def get_current_datetime() -> Dict[str, str]:
        """Get current date and time"""
        now = datetime.now()
        return {
            "datetime": now.strftime("%Y-%m-%d %H:%M:%S"),
            "date": now.strftime("%Y-%m-%d"),
            "time": now.strftime("%H:%M:%S"),
            "day": now.strftime("%A"),
            "timestamp": now.timestamp()
        }
    
    @staticmethod
    def get_formatted_time(format_str: str = "%Y-%m-%d %H:%M:%S") -> str:
        """Get formatted current time"""
        return datetime.now().strftime(format_str)


class ToolRouter:
    """Intelligently routes between available tools"""
    
    def __init__(self, serper_api_key: Optional[str] = None):
        self.web_search = WebSearchTool(serper_api_key) if serper_api_key else None
        self.doc_creator = DocumentCreationTool()
        self.datetime_tool = DateTimeTool()
    
    def route(self, query: str, tool_name: str, **kwargs) -> Any:
        """Route to appropriate tool based on request"""
        if tool_name == "web_search":
            if self.web_search:
                depth = kwargs.get("depth", "shallow")
                return self.web_search.search(query, depth)
            else:
                return {"error": "Serper API key not provided. Add it in the sidebar to enable web search."}
        
        elif tool_name == "datetime":
            return self.datetime_tool.get_current_datetime()
        
        elif tool_name == "create_pdf":
            content = kwargs.get("content", "")
            filename = kwargs.get("filename", "document.pdf")
            return self.doc_creator.create_pdf(content, filename)
        
        elif tool_name == "create_word":
            content = kwargs.get("content", "")
            filename = kwargs.get("filename", "document.docx")
            return self.doc_creator.create_word(content, filename)
        
        elif tool_name == "create_excel":
            data = kwargs.get("data", [])
            filename = kwargs.get("filename", "spreadsheet.xlsx")
            return self.doc_creator.create_excel(data, filename)
        
        else:
            return {"error": f"Unknown tool: {tool_name}"}
    
    def detect_tool_need(self, query: str) -> List[str]:
        """Detect which tools might be needed for a query"""
        query_lower = query.lower()
        needed_tools = []
        
        # Check for web search needs - be more aggressive
        web_keywords = ["search", "find", "look up", "latest", "current", "news", "google", 
                       "what is", "who is", "information about", "tell me about", "trends",
                       "recent", "update", "happening", "going on"]
        if any(word in query_lower for word in web_keywords):
            needed_tools.append("web_search")
        
        # Check for datetime needs
        time_keywords = ["time", "date", "today", "now", "when", "current day", 
                        "what day", "what time", "clock", "year", "month"]
        if any(word in query_lower for word in time_keywords):
            needed_tools.append("datetime")
        
        # Check for document creation needs
        if any(word in query_lower for word in ["pdf", "create document", "generate pdf", "make pdf"]):
            needed_tools.append("create_pdf")
        if any(word in query_lower for word in ["word", "docx", "word document", "create word"]):
            needed_tools.append("create_word")
        if any(word in query_lower for word in ["excel", "spreadsheet", "xlsx", "create excel"]):
            needed_tools.append("create_excel")
        
        return needed_tools
